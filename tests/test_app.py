import os
from pathlib import Path

os.environ["AFM_DATA_DIR"] = str(Path(__file__).parent / "runtime")

import app
from fastapi.testclient import TestClient


def test_full_pipeline(tmp_path, monkeypatch):
    def fake_available():
        return True

    def fake_json(endpoint, payload, timeout=90):
        if endpoint == "/api/embed":
            text = payload["input"].lower()
            dims = [0.0] * 8
            for word in text.split():
                dims[hash(word) % 8] += 1.0
            return {"embeddings": [dims]}
        if endpoint == "/api/generate":
            if "Return JSON only" in payload["prompt"]:
                return {"response": '{"title":"Renstra Proposal","doc_type":"proposal","summary":"Proposal for Renstra 2025-2029."}'}
            return {"response": "The matching document is the Renstra proposal."}
        return None

    monkeypatch.setattr(app, "ollama_available", fake_available)
    monkeypatch.setattr(app, "ollama_json", fake_json)

    sample = tmp_path / "mystery.docx"
    from docx import Document
    d = Document()
    d.add_heading("Proposal Renstra 2025-2029", level=1)
    d.add_paragraph("Dokumen ini membahas Renstra madrasah dan sasaran strategis.")
    d.save(sample)

    client = TestClient(app.app)

    r = client.post("/api/scan", json={"path": str(tmp_path)})
    assert r.status_code == 200, r.text
    assert r.json()["indexed"] == 1
    assert r.json()["errors"] == 0

    r = client.post("/api/search", json={"query": "Renstra", "limit": 10})
    assert r.status_code == 200
    assert len(r.json()["results"]) == 1

    r = client.post("/api/ask", json={"question": "Apa dokumen tentang Renstra?"})
    assert r.status_code == 200
    assert "Renstra" in r.json()["answer"]
    assert r.json()["sources"]
    assert r.json()["semantic"] is True

    doc_id = r.json()["sources"][0]["id"]
    r = client.post("/api/rename", json={"id": doc_id, "new_name": "2025_Renstra_Proposal.docx"})
    assert r.status_code == 200, r.text
    new_path = Path(r.json()["new_path"])
    assert new_path.name == "2025_Renstra_Proposal.docx"
    assert new_path.exists()


def test_search_does_not_require_all_question_words(tmp_path, monkeypatch):
    monkeypatch.setattr(app, "ollama_available", lambda: False)
    sample = tmp_path / "renstra.txt"
    sample.write_text("Renstra madrasah 2025-2029 sasaran strategis", encoding="utf-8")
    client = TestClient(app.app)
    r = client.post("/api/scan", json={"path": str(tmp_path)})
    assert r.status_code == 200
    r = client.post("/api/search", json={"query": "What documents mention Renstra?"})
    assert r.status_code == 200
    assert any(x["name"] == "renstra.txt" for x in r.json()["results"])
